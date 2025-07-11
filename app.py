import streamlit as st
import os
import io
import json
import re
import bcrypt
from datetime import datetime
import time
import base64 # Import base64 for image encoding

# --- Firebase Imports ---
import firebase_admin
from firebase_admin import credentials, auth, firestore
from firebase_admin import exceptions # Import exceptions module for FirebaseError

# --- Google Drive Imports ---
from google.oauth2 import service_account
from googleapiclient.discovery import build
from googleapiclient.http import MediaIoBaseUpload

# --- AI & Document Processing Imports ---
from openai import OpenAI
import pandas as pd
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn # For font color


# --- Streamlit Page Configuration (MUST BE THE FIRST ST COMMAND) ---
st.set_page_config(
    page_title="SSO Consultants AI Recruitment",
    page_icon="sso_logo.png", # Set favicon for the browser tab
    layout="wide" # Keeping wide layout, but centering content within it
)

# --- Custom CSS for Styling ---
st.markdown(
    """
    <style>
    /* Global base styling - pure white background, pure black text by default */
    body {
        background-color: #FFFFFF; /* Pure white background */
        color: #000000 !important; /* Pure black for general text readability - CRITICAL */
        font-family: 'Inter', sans-serif;
        height: 100vh; /* Ensure body takes full viewport height */
        margin: 0;
        padding: 0;
    }
    /* Streamlit's main app container */
    .stApp {
        background-color: #FFFFFF;
        color: #000000;
        min-height: 100vh; /* Ensure it takes full viewport height for centering */
        display: flex;
        flex-direction: column;
        justify-content: flex-start; /* Let content align to start, specific containers will center */
        align-items: stretch; /* Allow children to stretch, don't force overall app content to center here */
    }
    /* Headers */
    h1, h2, h3, h4, h5, h6 {
        color: #4CAF50; /* Green for headers from config.toml */
    }
    h1 {
        margin-top: 0px !important; /* Force no top margin */
        padding-top: 0px !important; /* Force no top padding */
    }
    /* Buttons */
    .stButton>button {
        background-color: #4CAF50; /* Green background for buttons from config.toml */
        color: white; /* White text for buttons */
        border-radius: 5px;
        border: none;
        padding: 0px 5px;
        font-size: 16px;
        cursor: pointer;
        transition: background-color 0.3s ease;
        
        /* New/modified properties for consistent sizing and centering text */
        width: 100%; /* Make button fill its immediate container (e.g., column) */
        height: 50px; /* Set a fixed ideal height for all buttons */
        display: flex; /* Use flexbox to center text vertically and horizontally */
        align-items: center;
        justify-content: center;
        text-align: center; /* Fallback for text alignment */
        white-space: normal; /* Allow text to wrap if it's too long for the button */
        word-break: break-word; /* Break words if necessary */
    }
    .stButton>button:hover {
        background-color: #66BB6A; /* Lighter green on hover */
    }
    /* Text Inputs, Text Areas, Select Boxes */
    .stTextInput>div>div>input, .stTextArea>div>div>textarea, .stSelectbox>div>div {
        border: 1px solid #4CAF50; /* Green border */
        border-radius: 5px;
        padding: 8px;
        color: #000000; /* Black text */
        background-color: #FFFFFF; /* White background */
    }
    /* Placeholder text */
    .stTextInput>div>div>input::placeholder, .stTextArea>div>div>textarea::placeholder {
        color: #888888; /* Grey placeholder */
    }
    /* Specific Streamlit components by data-testid */
    .stTabs [data-baseweb="tab-list"] button [data-testid="stMarkdownContainer"] p {
        font-size:1.2rem;
        color: #4CAF50; /* Green tab headers */
    }
    /* Centering content for specific pages */
    .main-content-centered {
        display: flex;
        flex-direction: column;
        align-items: center;
        text-align: center;
        padding-top: 0px;
    }
    .main-content-centered .stTextInput, .main-content-centered .stButton {
        max-width: 400px; /* Limit width of input/buttons in centered view */
        width: 100%;
    }

    /* New CSS for the fixed top-left logo */
    .fixed-top-left-logo {
        position: fixed;
        top: 15px; /* Distance from the top */
        left: 20px; /* Distance from the left */
        z-index: 1000; /* Ensure it stays on top of other elements */
        height: 60px; /* Adjust height as needed */
        width: auto;
    }

    /* Adjusted .login-container to remove min-height and border/shadow */
    .login-container {
        display: flex;
        flex-direction: column;
        align-items: center;
        width: 100%;
        max-width: 650px; /* MODIFIED: Increased width from 550px to 650px */
        margin: auto; /* Centers the content within its flex parent horizontally */
        padding: 20px; /* Reduced padding */
        background-color: #FFFFFF;
        /* No border, no box-shadow */
    }
    .login-container .stTextInput, .login-container .stButton {
        width: 100%; /* Ensure inputs and buttons fill the container width */
    }

    /* --- ALIGNMENT RULES --- */

    /* Target the main content area (excluding sidebar) */
    /* This targets the 'section' element with class 'main' that Streamlit uses for the central content */
    section.main {
        display: flex;
        flex-direction: column;
        align-items: center; /* Center content horizontally within the main section */
        flex-grow: 1; /* Allow it to grow and fill available vertical space */
        width: 100%; /* Ensure it takes full width */
        padding: 0 20px; /* Add some horizontal padding to prevent content from touching edges */
    }

    /* Streamlit's internal blocks within the main section */
    /* MODIFIED: Removed align-items: center; from here to allow more control */
    [data-testid="stVerticalBlock"] {
        display: flex;
        flex-direction: column;
        /* align-items: center;  REMOVED this line */
        width: 100%; /* Ensure it takes full width of its parent */
    }

    /* Adjust Streamlit's root element to allow flex centering */
    /* This centers the entire stApp component on the page */
    #root > div:first-child {
        display: flex;
        flex-direction: column;
        justify-content: center; /* Center vertically */
        align-items: center; /* Center horizontally */
        min-height: 100vh;
        width: 100%;
    }
    
    /* Ensure forms are centered if they are not inside a flex-centered container (like login-container) */
    form:not(.login-container form) { 
        width: 100%;
        max-width: 500px; /* Adjust as desired for general forms */
        margin: 0 auto; /* Center the form itself */
    }

    /* Sidebar content centering */
    [data-testid="stSidebarContent"] {
        display: flex;
        flex-direction: column;
        align-items: center; /* Centers content horizontally within the sidebar */
        padding-top: 20px; /* Adjust as needed */
    }
    /* For images specifically inside sidebar to ensure centering */
    [data-testid="stSidebarContent"] img {
        display: block; /* Important for margin: auto to work */
        margin-left: auto;
        margin-right: auto;
    }

    /* --- NEW CSS FOR LEFT-ALIGNED FILE UPLOADERS --- */
    .left-aligned-content {
        width: 100%; /* Ensure it takes full width of its parent */
        display: flex;
        flex-direction: column;
        align-items: flex-start; /* Aligns children to the left */
    }
    /* Ensure file uploader widgets themselves also align left within this container */
    .left-aligned-content [data-testid="stFileUploader"] {
        width: 100%; /* Take full width of the left-aligned container */
    }
    /* This targets the label of the file uploader */
    .left-aligned-content [data-testid="stFileUploader"] label {
        text-align: left !important; /* Force label text to left align */
        width: 100%; /* Ensure label spans full width */
    }
    /* This targets the inner vertical block within the file uploader */
    .left-aligned-content [data-testid="stFileUploader"] > div > [data-testid="stVerticalBlock"] {
        align-items: flex-start !important; /* Force inner vertical blocks also align left - ADDED !important */
    }
    /* This specifically targets the upload button text within the file uploader */
    .left-aligned-content [data-testid="stFileUploader"] button div span {
        text-align: left !important;
        width: 100%;
        justify-content: flex-start; /* Align button content to start */
    }
    /* Ensure the button itself also aligns left within the uploader */
    .left-aligned-content [data-testid="stFileUploader"] button {
        align-self: flex-start; /* Align the button element itself to the left */
    }

    /* --- Fix for Streamlit selectbox dropdown height and text wrapping --- */
    /* Target the popover containing the selectbox options */
    div[data-baseweb="select-popover"] {
        max-height: 300px; /* Set a maximum height for the dropdown */
        overflow-y: auto; /* Enable vertical scrolling if content exceeds max-height */
    }

    /* Ensure text in selectbox options wraps and height adjusts */
    div[data-baseweb="select-popover"] div[role="option"] {
        white-space: normal !important; /* Allow text to wrap within each option */
        height: auto !important; /* Ensure height adjusts to content */
        min-height: 38px; /* Set a minimum height for readability if content is short */
        display: flex; /* Use flex to align content if needed */
        align-items: center; /* Center vertically if single line */
        padding: 8px 12px; /* Add some padding for better appearance */
    }

    /* START OF NEW CODE TO FIX SELECTBOX WIDTH */
    [data-testid="stSelectbox"] > div[data-baseweb="select"] > div[role="button"] {
        min-width: 450px !important; /* Set a generous minimum width for the display area */
        max-width: none !important; /* Ensure no max-width is constraining it */
        width: 100% !important; /* Try to make it take full available width */
    }
    /* For the dropdown options themselves to also have enough width */
    [data-testid="stSelectbox"] > div[data-baseweb="popover"] {
        min-width: 450px !important; /* Ensure dropdown options also have enough width */
        max-width: none !important; /* Ensure options are not constrained */
    }
    /* END OF NEW CODE TO FIX SELECTBOX WIDTH */

    </style>
    """,
    unsafe_allow_html=True
)

# --- Session State Initialization ---
if 'logged_in' not in st.session_state:
    st.session_state['logged_in'] = False
if 'user_email' not in st.session_state:
    st.session_state['user_email'] = None
if 'user_uid' not in st.session_state:
    st.session_state['user_uid'] = None
if 'is_admin' not in st.session_state:
    st.session_state['is_admin'] = False
if 'login_mode' not in st.session_state: # New: Control login flow
    st.session_state['login_mode'] = 'choose_role' # Start with role selection
if 'is_admin_attempt' not in st.session_state: # New: Flag for current login attempt type
    st.session_state['is_admin_attempt'] = False # Default to user login attempt
if 'username' not in st.session_state:
    st.session_state['username'] = None
if 'has_set_username' not in st.session_state:
    st.session_state['has_set_username'] = False
if 'needs_username_setup' not in st.session_state: # Flag to explicitly trigger setup page
    st.session_state['needs_username_setup'] = False
if 'login_success' not in st.session_state:
    st.session_state['login_success'] = False
if 'current_admin_page' not in st.session_state:
    st.session_state['current_admin_page'] = 'reports'


# --- Firebase Initialization ---
# Ensure only one app instance is initialized
if not firebase_admin._apps:
    try:
        # Load the Firebase service account JSON string from Streamlit secrets
        # The key name is "FIREBASE_SERVICE_ACCOUNT_KEY" as per your secrets.toml
        firebase_service_account_json_str = st.secrets["FIREBASE_SERVICE_ACCOUNT_KEY"]
        # Parse the JSON string into a Python dictionary
        firebase_service_account_info = json.loads(firebase_service_account_json_str)

        # Explicitly replace '\n' with actual newline characters in the private_key
        # This is crucial because Streamlit's TOML parser might preserve '\\n' as literal escapes
        # if the original JSON file had them, and firebase_admin expects actual newlines.
        if "private_key" in firebase_service_account_info:
            firebase_service_account_info["private_key"] = firebase_service_account_info["private_key"].replace('\\n', '\n')

        cred = credentials.Certificate(firebase_service_account_info)
        firebase_admin.initialize_app(cred) # Use firebase_admin.initialize_app
        db = firestore.client()
        st.success("Firebase initialized successfully.")

    except KeyError:
        st.error("Firebase 'FIREBASE_SERVICE_ACCOUNT_KEY' not found in Streamlit secrets! "
                 "Please add your Firebase service account JSON content as a multi-line string "
                 "to your app's secrets.toml under the key 'FIREBASE_SERVICE_ACCOUNT_KEY'.")
        st.stop()
    except json.JSONDecodeError as e:
        st.error(f"Error decoding Firebase service account JSON: {e}. "
                 "Please ensure 'FIREBASE_SERVICE_ACCOUNT_KEY' in secrets.toml is valid JSON.")
        st.stop()
    except Exception as e:
        st.error(f"An unexpected error occurred during Firebase initialization: {e}")
        st.info("Please ensure your 'FIREBASE_SERVICE_ACCOUNT_KEY' is valid and correctly formatted in secrets.toml.")
        st.stop()
else:
    db = firestore.client()

# --- Google Drive Configuration ---
drive_service = None
try:
    # Load the Google Drive key JSON string from Streamlit secrets
    # The key name is "GOOGLE_DRIVE_KEY" as per your secrets.toml
    google_drive_key_json_str = st.secrets["GOOGLE_DRIVE_KEY"]
    # Parse the JSON string into a Python dictionary
    google_drive_key_info = json.loads(google_drive_key_json_str)

    # Explicitly replace '\n' with actual newline characters in the private_key
    if "private_key" in google_drive_key_info:
        google_drive_key_info["private_key"] = google_drive_key_info["private_key"].replace('\\n', '\n')

    # Define the necessary scopes for Google Drive access
    SCOPES = ['https://www.googleapis.com/auth/drive'] # Scope for full Drive access

    # Create credentials from the service account info
    drive_credentials = service_account.Credentials.from_service_account_info(google_drive_key_info, scopes=SCOPES)

    # Build the Google Drive API service client
    drive_service = build('drive', 'v3', credentials=drive_credentials)
    

except KeyError:
    st.error("Google Drive 'GOOGLE_DRIVE_KEY' not found in Streamlit secrets! "
             "Please add your Google Drive service account JSON content as a multi-line string "
             "to your app's secrets.toml under the key 'GOOGLE_DRIVE_KEY'.")
    st.stop()
except json.JSONDecodeError as e:
    st.error(f"Error decoding Google Drive key JSON: {e}. "
             "Please ensure 'GOOGLE_DRIVE_KEY' in secrets.toml is valid JSON.")
    st.stop()
except Exception as e:
    st.error(f"An unexpected error occurred during Google Drive initialization: {e}")
    st.info("Please ensure your 'GOOGLE_DRIVE_KEY' is valid and Google Drive API is enabled in secrets.toml.")
    st.stop()

# --- OpenAI API Key Setup ---
openai_client = None
try:
    # Load the OpenAI API key from Streamlit secrets
    openai_api_key = st.secrets["OPENAI_API_KEY"]
    # Initialize the OpenAI client with the loaded API key
    openai_client = OpenAI(api_key=openai_api_key)
    

except KeyError:
    st.error("OPENAI_API_KEY not found in Streamlit secrets! "
             "Please add your OpenAI API key to your app's secrets "
             "on Streamlit Community Cloud under the key 'OPENAI_API_KEY'.")
    st.stop()
except Exception as e:
    st.error(f"An unexpected error occurred during OpenAI API key setup: {e}")
    st.stop()

# --- Google Drive Reports Folder ID (from secrets) ---
GOOGLE_DRIVE_REPORTS_FOLDER_ID = None
try:
    GOOGLE_DRIVE_REPORTS_FOLDER_ID = st.secrets["GOOGLE_DRIVE_REPORTS_FOLDER_ID"]
except KeyError:
    st.error("GOOGLE_DRIVE_REPORTS_FOLDER_ID not found in Streamlit secrets! "
             "Please add the ID of your Google Drive reports folder to your app's secrets "
             "on Streamlit Community Cloud under the key 'GOOGLE_DRIVE_REPORTS_FOLDER_ID'.")
    st.stop()


# --- Utility Functions ---

# Function to hash passwords for Firestore storage (for local emulator login)
def hash_password(password):
    return bcrypt.hashpw(password.encode('utf-8'), bcrypt.gensalt()).decode('utf-8')

# Function to check password against hash (for local emulator login)
def check_password(password, hashed_password):
    return bcrypt.checkpw(password.encode('utf-8'), hashed_password.encode('utf-8'))

def login_user(email, password): # Removed desired_login_type, as is_admin is from DB
    try:
        user = auth.get_user_by_email(email)
        user_doc_ref = db.collection('users').document(user.uid)
        user_doc = user_doc_ref.get()

        if user_doc.exists:
            user_data = user_doc.to_dict()
            hashed_password_from_db = user_data.get('hashed_password')
            is_admin_from_db = user_data.get('is_admin', False)
            username_from_db = user_data.get('username')
            has_set_username_from_db = user_data.get('has_set_username', False)

            if not hashed_password_from_db or not check_password(password, hashed_password_from_db):
                 st.error("Invalid credentials. Please check your password.")
                 return

            # Check if the attempted login type matches the user's actual role
            if st.session_state['is_admin_attempt'] and not is_admin_from_db:
                st.error("You attempted to log in as an Admin, but this account does not have admin privileges.")
                return
            if not st.session_state['is_admin_attempt'] and is_admin_from_db:
                st.warning("You logged in as a User, but this account has admin privileges. You can log in as Admin to access more features.")
                # Allow login but keep the current attempt as user. This might be desired.
                # If strict, would be: st.error("Please log in as Admin for this account.")
                # return

            st.session_state['logged_in'] = True
            st.session_state['user_email'] = email
            st.session_state['user_uid'] = user.uid
            st.session_state['is_admin'] = is_admin_from_db # Actual role from DB
            st.session_state['username'] = username_from_db
            st.session_state['has_set_username'] = has_set_username_from_db
            st.session_state['login_mode'] = 'logged_in' # Transition to logged-in state

            # Check if user needs to set username/password
            if not st.session_state['has_set_username']: 
                st.session_state['needs_username_setup'] = True
                st.success(f"Welcome, {email}! Please set up your display name and password.")
            else:
                st.session_state['needs_username_setup'] = False # Ensure this is false if already set
                if st.session_state['username']:
                    st.success(f"Logged in as {st.session_state['username']}!")
                else:
                    st.success(f"Logged in as {st.session_state['user_email']}!") # Fallback if username not set
            st.session_state['login_success'] = True
            
            if st.session_state['is_admin']:
                st.session_state['current_admin_page'] = 'generate' # Admin default to Generate Report
            else:
                st.session_state['current_admin_page'] = 'generate' # User default to Generate Report (MODIFIED)

            st.rerun() # Rerun to update UI after login
        else:
            st.error("User not found or password incorrect.")
    except auth.UserNotFoundError:
        st.error("User not found.")
    except Exception as e:
        st.error(f"Login error: {e}")

def create_user(email, password, is_admin=False):
    try:
        user_record = auth.create_user(email=email, password=password)
        user_ref = db.collection('users').document(user_record.uid)
        user_ref.set({
            'email': email,
            'is_admin': is_admin,
            'created_at': firestore.SERVER_TIMESTAMP,
            'hashed_password': hash_password(password), # Store hash for local emulator login
            'username': None, # New field: initially none
            'has_set_username': False # New flag: user needs to set username and password
        })
        st.success(f"User {email} created successfully!")
        return user_record.uid
    except exceptions.FirebaseError as e:
        error_message = e.code
        if "email-already-exists" in error_message:
            st.error("The email address is already in use by another account.")
        else:
            st.error(f"Error creating user: {error_message}")
        return None
    except Exception as e:
        st.error(f"An unexpected error occurred: {e}")
        return None

def logout_user():
    for key in ['logged_in', 'user_email', 'user_uid', 'is_admin', 'username', 'has_set_username', 'needs_username_setup', 'login_success', 'current_admin_page', 'login_mode', 'is_admin_attempt']:
        if key in st.session_state:
            del st.session_state[key]
    st.session_state['login_mode'] = 'choose_role' # Reset to role selection on logout
    st.rerun()

# --- Content Extraction Functions ---
def get_pdf_text(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def get_docx_text(file):
    document = Document(file)
    text = ""
    for paragraph in document.paragraphs:
        text += paragraph.text + "\n"
    return text

# --- OpenAI/AI Functions ---
def get_openai_response(prompt_text, json_mode=False):
    # Use the globally initialized openai_client
    if openai_client:
        try:
            messages = [
                {"role": "system", "content": "You are a helpful AI assistant specialized in analyzing Job Descriptions and CVs. Provide concise, direct, and actionable insights. Be professional and objective."},
                {"role": "user", "content": prompt_text}
            ]
            
            if json_mode:
                response = openai_client.chat.completions.create(
                    model="gpt-4o",
                    messages=messages,
                    response_format={ "type": "json_object" }, # Enable JSON mode
                    temperature=0.7
                )
                return json.loads(response.choices[0].message.content) # Parse JSON
            else:
                response = openai_client.chat.completions.create(
                    model="gpt-4o", # Using a powerful model
                    messages=messages,
                    temperature=0.7 # Adjust creativity
                )
                return response.choices[0].message.content
        except Exception as e:
            # Add a print statement to ensure it goes to console logs
            print(f"DEBUG: Caught error in get_openai_response: Type={type(e).__name__}, Message={e}")
            st.error(f"Error calling OpenAI API: {e}. Please check your API key and network connection. If the error persists, try reducing the complexity of the prompt or input files.")
            return "Error: Could not get response from AI." if not json_mode else {"error": "Could not get response from AI."}
    else:
        st.error("OpenAI client not initialized. Cannot generate AI response.")
        return "Error: OpenAI client not available." if not json_mode else {"error": "OpenAI client not available."}


# --- NEW AI PROMPT HELPER FUNCTIONS FOR STRUCTURED DATA ---

def get_candidate_evaluation_data(jd_text, cv_texts, cv_filenames):
    evaluations = []
    for i, cv_text in enumerate(cv_texts):
        prompt = f"""
        Given the following Job Description (JD) and Candidate CV, evaluate the candidate and provide the following details in a JSON object:
        - CandidateName: Full name of the candidate (deduce from CV).
        - MatchPercent: An integer percentage (e.g., 75) indicating overall match with the JD.
        - Ranking: An integer rank (e.g., 1, 2, 3) relative to other candidates, assuming this is the only candidate evaluated right now. Assign rank 1.
        - ShortlistProbability: "High", "Moderate", or "Low".
        - KeyStrengths: A concise string listing key strengths of the CV relative to the JD.
        - KeyGaps: A concise string listing key areas of improvement/gaps in the CV relative to the JD.
        - LocationSuitability: "Suitable", "Consider", or "Not Suitable" (based on JD's location if specified, and CV's implied location).
        - Comments: A concise overall comment on the candidate's fit.

        Job Description:
        {jd_text}

        Candidate CV ({cv_filenames[i]}):
        {cv_text}

        Ensure the output is a valid JSON object.
        """
        response = get_openai_response(prompt, json_mode=True)
        if isinstance(response, dict) and "error" not in response:
            # Add filename for internal tracking
            response['OriginalFilename'] = cv_filenames[i]
            evaluations.append(response)
        else:
            st.warning(f"Could not get structured evaluation for {cv_filenames[i]}: {response.get('error', 'Unknown error')}")
            evaluations.append({
                "CandidateName": cv_filenames[i].replace(".pdf", "").replace(".docx", ""),
                "MatchPercent": 0,
                "Ranking": 99,
                "ShortlistProbability": "Low",
                "KeyStrengths": "AI analysis failed.",
                "KeyGaps": "AI analysis failed.",
                "LocationSuitability": "Unknown",
                "Comments": "Failed to generate AI analysis."
            })
    
    # After getting individual evaluations, re-rank them globally based on MatchPercent
    if evaluations:
        evaluations.sort(key=lambda x: x.get('MatchPercent', 0), reverse=True)
        for rank, eval_data in enumerate(evaluations):
            eval_data['Ranking'] = rank + 1
            # Adjust ShortlistProbability based on sorted rank for multi-candidate view
            if eval_data['MatchPercent'] >= 80:
                eval_data['ShortlistProbability'] = "High"
            elif eval_data['MatchPercent'] >= 60:
                eval_data['ShortlistProbability'] = "Moderate"
            else:
                eval_data['ShortlistProbability'] = "Low"
    
    return evaluations


def get_criteria_comparison_data(jd_text, cv_texts, cv_filenames, criteria_list):
    # This prompt asks the AI to evaluate each candidate against a fixed set of criteria
    # and provide a simple emoji-based rating.
    prompt = f"""
    Given the Job Description and the following CVs, evaluate each candidate against the provided criteria.
    For each candidate and each criterion, provide an emoji:
    - ✅ for strong match/presence
    - ⚠️ for partial match/some presence/needs consideration
    - ❌ for no match/significant gap
    Output should be a JSON object where keys are the criteria and values are objects containing candidate names as keys and their emoji ratings as values.

    Job Description:
    {jd_text}

    Candidate CVs:
    """
    for i, cv_text in enumerate(cv_texts):
        prompt += f"\n--- CV {cv_filenames[i]} ---\n{cv_text}\n"

    prompt += f"\nCriteria to evaluate (use these exact names as keys): {', '.join(criteria_list)}"
    prompt += "\nExample JSON structure: {'Education (MBA)': {'Candidate1 Name': '✅', 'Candidate2 Name': '⚠️'}, 'Relevant Experience': {'Candidate1 Name': '❌', 'Candidate2 Name': '✅'}}"

    response = get_openai_response(prompt, json_mode=True)
    if isinstance(response, dict) and "error" not in response:
        return response
    else:
        st.warning(f"Could not get structured criteria comparison: {response.get('error', 'Unknown error')}")
        return {criterion: {filename.replace('.pdf','').replace('.docx',''): "❌" for filename in cv_filenames} for criterion in criteria_list} # Fallback

def get_general_observations_and_shortlist(evaluations):
    # Sort candidates by ranking to feed into the prompt correctly
    sorted_candidates = sorted(evaluations, key=lambda x: x.get('Ranking', 99))

    prompt = "Based on the following candidate evaluations, provide:\n"
    prompt += "1. General Observations: An overall summary of the candidate pool, highlighting top candidates and general trends.\n"
    prompt += "2. Final Shortlist Recommendation: A list of names of candidates recommended for shortlisting, based primarily on 'High' or 'Moderate' shortlist probability and ranking.\n\n"
    prompt += "Candidate Evaluations (sorted by rank):\n"
    for cand in sorted_candidates:
        prompt += f"- {cand['CandidateName']} (Match: {cand['MatchPercent']}%, Rank: {cand['Ranking']}, Shortlist: {cand['ShortlistProbability']}): Strengths: {cand['KeyStrengths']}. Gaps: {cand['KeyGaps']}. Comments: {cand['Comments']}\n"
    
    prompt += "\nOutput in JSON format with keys 'GeneralObservations' (string) and 'ShortlistedCandidates' (list of strings)."

    response = get_openai_response(prompt, json_mode=True)
    if isinstance(response, dict) and "error" not in response:
        return response
    else:
        st.warning(f"Could not get general observations and shortlist: {response.get('error', 'Unknown error')}")
        return {"GeneralObservations": "Could not generate general observations.", "ShortlistedCandidates": []}


# --- Report Generation Function (MODIFIED FOR NEW FORMAT) ---
def create_comparative_docx_report(jd_text, cv_texts, report_data, candidate_evaluations, criteria_comparison_data, general_and_shortlist_data):
    document = Document()

    document.add_heading('JD-CV Comparative Analysis Report', level=1)
    
    # Add a paragraph for general info
    document.add_paragraph(f"Generated by {report_data.get('generated_by_username', report_data.get('generated_by_email'))}")
    document.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    document.add_paragraph(f"Job Description: {report_data.get('jd_filename', 'N/A')}")
    document.add_paragraph(f"Candidates: {', '.join(report_data.get('cv_filenames', ['N/A']))}")
    
    document.add_page_break()

    # --- Candidate Evaluation Table ---
    document.add_heading('🧾 Candidate Evaluation Table', level=2)
    document.add_paragraph('Detailed assessment of each candidate against the Job Description:')

    if candidate_evaluations:
        headers = ["Candidate Name", "Match %", "Ranking", "Shortlist Probability", "Key Strengths", "Key Gaps", "Location Suitability", "Comments"]
        table = document.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'

        # Add header row
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(headers):
            hdr_cells[i].text = header_text
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # Add data rows
        for candidate in candidate_evaluations:
            row_cells = table.add_row().cells
            row_cells[0].text = candidate.get('CandidateName', 'N/A')
            row_cells[1].text = f"{candidate.get('MatchPercent', 0)}%"
            row_cells[2].text = str(candidate.get('Ranking', 'N/A'))
            row_cells[3].text = candidate.get('ShortlistProbability', 'N/A')
            row_cells[4].text = candidate.get('KeyStrengths', 'N/A')
            row_cells[5].text = candidate.get('KeyGaps', 'N/A')
            row_cells[6].text = candidate.get('LocationSuitability', 'N/A')
            row_cells[7].text = candidate.get('Comments', 'N/A')
    else:
        document.add_paragraph("No candidate evaluation data available.")

    document.add_page_break()

    # --- Criteria Comparison Table ---
    document.add_heading('✅ Additional Observations (Criteria Comparison)', level=2)

    if criteria_comparison_data and candidate_evaluations:
        # Get all unique candidate names from evaluations to ensure consistent column order
        candidate_names_ordered = [cand['CandidateName'] for cand in candidate_evaluations]
        
        # Prepare headers: "Criteria" + all candidate names
        criteria_headers = ["Criteria"] + candidate_names_ordered
        
        # Determine number of rows (number of criteria)
        num_criteria = len(criteria_comparison_data)
        
        table = document.add_table(rows=num_criteria + 1, cols=len(criteria_headers))
        table.style = 'Table Grid'

        # Add header row for criteria comparison
        hdr_cells = table.rows[0].cells
        for i, header_text in enumerate(criteria_headers): # Use criteria_headers here
            hdr_cells[i].text = header_text
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
            hdr_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        
        # Add data rows
        row_idx = 1
        for criteria, candidate_ratings in criteria_comparison_data.items():
            row_cells = table.rows[row_idx].cells
            row_cells[0].text = criteria # First cell is the criterion name
            row_cells[0].paragraphs[0].runs[0].font.bold = True # Bold the criterion name

            for col_idx, cand_name in enumerate(candidate_names_ordered):
                # Use .get() with a default for robustness
                emoji = candidate_ratings.get(cand_name, 'N/A')
                row_cells[col_idx + 1].text = emoji # +1 because first col is 'Criteria'
                row_cells[col_idx + 1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER # Center emojis
            row_idx += 1
    else:
        document.add_paragraph("No criteria comparison data available.")

    document.add_page_break()

    # --- General Observations and Shortlist ---
    document.add_heading('General Observations', level=2)
    document.add_paragraph(general_and_shortlist_data.get('GeneralObservations', 'No general observations available.'))

    document.add_heading('📌 Final Shortlist Recommendation', level=2)
    if general_and_shortlist_data.get('ShortlistedCandidates'):
        document.add_paragraph(f"Shortlisted candidates: {', '.join(general_and_shortlist_data.get('ShortlistedCandidates'))}")
    else:
        document.add_paragraph("No candidates recommended for shortlist based on current analysis.")

    # Save the document to a BytesIO object
    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# --- Pages/UI Functions ---

def display_login_form():
    st.image("sso_logo.png", width=100)
    st.markdown("<h2 style='text-align: center; color: #4CAF50;'>SSO Consultants AI Recruitment</h2>", unsafe_allow_html=True)
    st.markdown("<p style='text-align: center;'>Please choose your login type.</p>", unsafe_allow_html=True)

    col_user, col_admin = st.columns(2)
    with col_user:
        if st.button("Login as User", key="user_role_button"):
            st.session_state['login_mode'] = 'login_form'
            st.session_state['is_admin_attempt'] = False
            st.rerun()
    with col_admin:
        if st.button("Login as Admin", key="admin_role_button"):
            st.session_state['login_mode'] = 'login_form'
            st.session_state['is_admin_attempt'] = True
            st.rerun()

def show_login_and_create_account_forms():
    st.image("sso_logo.png", width=100) # Logo on login page
    st.markdown(f"<h2 style='text-align: center; color: #4CAF50;'>Login {'(Admin)' if st.session_state.get('is_admin_attempt') else '(User)'}</h2>", unsafe_allow_html=True)
    
    with st.form("login_form"):
        st.markdown("<p style='text-align: center;'>Already have an account? Log in here:</p>", unsafe_allow_html=True)
        email = st.text_input("Email", key="login_email")
        password = st.text_input("Password", type="password", key="login_password")
        submitted = st.form_submit_button("Login")
        if submitted:
            login_user(email, password)

    if st.button("Back to Role Selection", key="back_to_role_selection_button"):
        st.session_state['login_mode'] = 'choose_role'
        st.rerun()

def setup_username_and_password_page():
    st.image("sso_logo.png", width=100)
    st.markdown("<h2 style='text-align: center; color: #4CAF50;'>Complete Your Profile</h2>", unsafe_allow_html=True)
    st.write("Welcome! Please set a display name and your new password.")

    with st.form("username_setup_form"):
        new_username = st.text_input("Display Name (e.g., Your Name)", value=st.session_state.get('username', ''), key="setup_username")
        new_password = st.text_input("New Password", type="password", key="setup_new_password")
        confirm_password = st.text_input("Confirm New Password", type="password", key="setup_confirm_password")
        
        submitted = st.form_submit_button("Save Profile")

        if submitted:
            if not new_username:
                st.error("Display Name cannot be empty.")
            elif new_password != confirm_password:
                st.error("Passwords do not match.")
            elif len(new_password) < 6:
                st.error("Password must be at least 6 characters long.")
            else:
                try:
                    # Update password in Firebase Auth
                    auth.update_user(st.session_state['user_uid'], password=new_password)
                    
                    # Update username and hashed password in Firestore
                    user_doc_ref = db.collection('users').document(st.session_state['user_uid'])
                    user_doc_ref.update({
                        'username': new_username,
                        'hashed_password': hash_password(new_password),
                        'has_set_username': True
                    })
                    st.session_state['username'] = new_username
                    st.session_state['has_set_username'] = True
                    st.session_state['needs_username_setup'] = False # Profile setup complete
                    st.success("Profile updated successfully!")
                    # --- NEW: Explicitly log out and prompt for re-login ---
                    st.info("Please log in again with your new password.")
                    logout_user() # This will rerun and go to choose_role/login page
                    # --- END NEW ---
                except exceptions.FirebaseError as e:
                    st.error(f"Error updating profile: {e}")
                except Exception as e:
                    st.error(f"An unexpected error occurred: {e}")

def generate_comparative_report_page():
    # Centered Title (Replaced st.title with markdown for more control)
    st.markdown("<h1 style='text-align: center; color: #4CAF50;'>SSO Consultants AI Recruitment Tool</h1>", unsafe_allow_html=True)

    st.write(f"Logged in as: **{st.session_state.get('username', st.session_state.get('user_email', 'Guest'))}** {'(Admin)' if st.session_state['is_admin'] else ''}")
    st.subheader("Generate Comparative Analysis Report")

    st.write("Upload a Job Description (JD) and multiple Candidate CVs to generate a comparative analysis report.")

    jd_file = st.file_uploader("Upload Job Description (PDF/DOCX)", type=["pdf", "docx"], key="jd_uploader")
    cv_files = st.file_uploader("Upload Candidate CVs (PDF/DOCX)", type=["pdf", "docx"], accept_multiple_files=True, key="cv_uploader")

    # Define the list of criteria for comparison
    comparison_criteria = [
        "Education ", "Relevant Experience",  "Certifications",
        "Location Suitability", "Technical Skills", "Soft Skills"
    ]

    with st.expander("AI Report Generation Settings"):
        st.write("Customize the criteria the AI will use for comparison.")
        selected_criteria = []
        for criterion in comparison_criteria:
            if st.checkbox(criterion, value=True, key=f"criterion_{criterion.replace(' ', '_')}"):
                selected_criteria.append(criterion)
        
        if not selected_criteria:
            st.warning("Please select at least one criterion for comparison.")
            return # Prevent generation if no criteria are selected

    if st.button("Generate Report", key="generate_report_button"):
        if jd_file and cv_files:
            with st.spinner("Analyzing documents and generating report... This may take a few moments."):
                jd_text = ""
                if jd_file.type == "application/pdf":
                    jd_text = get_pdf_text(jd_file)
                elif jd_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                    jd_text = get_docx_text(jd_file)
                else:
                    st.error("Unsupported JD file type.")
                    return

                cv_texts = []
                cv_filenames = []
                for cv_file in cv_files:
                    cv_filenames.append(cv_file.name)
                    if cv_file.type == "application/pdf":
                        cv_texts.append(get_pdf_text(cv_file))
                    elif cv_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                        cv_texts.append(get_docx_text(cv_file))
                    else:
                        st.warning(f"Skipping unsupported CV file type: {cv_file.name}")
                        continue
                
                if not cv_texts:
                    st.error("No supported CV files found to analyze.")
                    return

                # Step 1: Get individual candidate evaluations
                st.info("Step 1/3: Evaluating individual candidates...")
                candidate_evaluations = get_candidate_evaluation_data(jd_text, cv_texts, cv_filenames)
                if any("Error: Could not get response from AI." in str(c.values()) for c in candidate_evaluations):
                    st.error("Failed to get complete candidate evaluations from AI. Report generation aborted.")
                    return

                # Step 2: Get criteria comparison data
                st.info("Step 2/3: Comparing candidates based on selected criteria...")
                criteria_comparison_data = get_criteria_comparison_data(jd_text, cv_texts, cv_filenames, selected_criteria)
                if any("error" in str(criteria_comparison_data.values()) for c in criteria_comparison_data.values()): # Check for errors in inner dicts
                    st.error("Failed to get criteria comparison from AI. Report generation aborted.")
                    return

                # Step 3: Get general observations and shortlist
                st.info("Step 3/3: Generating general observations and shortlist...")
                general_and_shortlist_data = get_general_observations_and_shortlist(candidate_evaluations)
                if "error" in general_and_shortlist_data.get('GeneralObservations', '').lower():
                    st.error("Failed to get general observations/shortlist from AI. Report generation aborted.")
                    return

                # Prepare report data for DOCX generation and Firestore
                report_data = {
                    "jd_filename": jd_file.name,
                    "cv_filenames": cv_filenames,
                    "generated_by_email": st.session_state['user_email'],
                    "generated_by_username": st.session_state['username'],
                    "timestamp": datetime.now().isoformat(), # ISO format for easy sorting in Firestore
                    "candidate_evaluations": candidate_evaluations,
                    "criteria_comparison_data": criteria_comparison_data,
                    "general_and_shortlist_data": general_and_shortlist_data
                }

                # Generate the DOCX report
                report_buffer = create_comparative_docx_report(
                    jd_text, cv_texts, report_data,
                    candidate_evaluations, criteria_comparison_data, general_and_shortlist_data
                )

                # Generate unique filename for the report
                username = report_data.get('generated_by_username', 'UnknownUser')
                timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
                report_full_filename = f"{username}_JD_CV_Analysis_Report_{timestamp}.docx" # MODIFIED LINE

                # Upload to Google Drive
                try:
                    if drive_service and GOOGLE_DRIVE_REPORTS_FOLDER_ID:
                        file_metadata = {
                            'name': report_full_filename,
                            'parents': [GOOGLE_DRIVE_REPORTS_FOLDER_ID],
                            'mimeType': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                        }
                        media = MediaIoBaseUpload(report_buffer, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', resumable=True)
                        uploaded_file = drive_service.files().create(body=file_metadata, media_body=media, fields='id').execute()
                        report_data['drive_file_id'] = uploaded_file.get('id')
                        st.success(f"Report uploaded to Google Drive: {report_full_filename}")
                except Exception as e:
                    st.error(f"Error uploading to Google Drive: {e}. You can still download the report directly.")
                    report_data['drive_file_id'] = None # Ensure it's marked as failed upload

                # Save report metadata to Firestore
                try:
                    reports_collection = db.collection('reports')
                    reports_collection.add(report_data)
                    st.success("Report metadata saved to database.")
                except Exception as e:
                    st.error(f"Error saving report metadata to database: {e}")

                st.success("Report generated and saved!")

                # Provide download link
                st.download_button(
                    label="Download Report",
                    data=report_buffer.getvalue(), # Use getvalue() after seeking to 0
                    file_name=report_full_filename,
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key="download_button"
                )
        else:
            st.error("Please upload both a Job Description and at least one CV to generate a report.")

def show_all_reports_page():
    st.markdown("<h1 style='text-align: center; color: #4CAF50;'>SSO Consultants AI Recruitment Tool</h1>", unsafe_allow_html=True)
    st.subheader("All Generated Reports")
    st.write(f"Logged in as: **{st.session_state.get('username', st.session_state.get('user_email', 'Guest'))}** {'(Admin)' if st.session_state['is_admin'] else ''}")

    reports_ref = db.collection('reports')
    if not st.session_state['is_admin']:
        # Filter reports by the current user's email if not admin
        reports_query = reports_ref.where('generated_by_email', '==', st.session_state['user_email']).order_by('timestamp', direction=firestore.Query.DESCENDING)
    else:
        # Admins can see all reports
        reports_query = reports_ref.order_by('timestamp', direction=firestore.Query.DESCENDING)

    try:
        reports_docs = reports_query.stream()
        reports = []
        for doc in reports_docs:
            report_data = doc.to_dict()
            reports.append({
                'id': doc.id,
                'jd_filename': report_data.get('jd_filename', 'N/A'),
                'cv_filenames': ', '.join(report_data.get('cv_filenames', ['N/A'])),
                'generated_by_email': report_data.get('generated_by_email', 'N/A'),
                'generated_by_username': report_data.get('generated_by_username', 'N/A'),
                'timestamp': datetime.fromisoformat(report_data['timestamp']).strftime('%Y-%m-%d %H:%M:%S') if 'timestamp' in report_data else 'N/A',
                'drive_file_id': report_data.get('drive_file_id'),
                'raw_data': report_data # Keep raw data for re-creating report if needed
            })

        if not reports:
            st.info("No reports found. Generate one in the 'Generate Report' section.")
            return

        st.dataframe(
            reports,
            column_order=["timestamp", "generated_by_username", "jd_filename", "cv_filenames"],
            hide_index=True,
            column_config={
                "timestamp": st.column_config.DatetimeColumn("Date & Time", format="YYYY-MM-DD HH:mm:ss"),
                "generated_by_username": "Generated By",
                "jd_filename": "Job Description",
                "cv_filenames": "CVs Analyzed",
                "id": None, # Hide internal ID
                "drive_file_id": None, # Hide drive ID in table
                "raw_data": None # Hide raw data in table
            },
            use_container_width=True
        )

        st.markdown("---")
        st.subheader("Actions on Reports")
        
        # Allow users to select a report from the displayed list
        selected_report_id = st.selectbox(
            "Select a report to view/download:",
            options=[r['id'] for r in reports], 
            format_func=lambda x: f"Report {next((r['timestamp'] for r in reports if r['id'] == x), 'N/A')} - {next((r['jd_filename'] for r in reports if r['id'] == x), 'N/A')}", 
            key="report_selector"
        )
        
        selected_report = next((r for r in reports if r['id'] == selected_report_id), None)

        if selected_report:
            if selected_report['drive_file_id']:
                # Generate a shareable link from Google Drive file ID
                drive_link = f"https://drive.google.com/file/d/{selected_report['drive_file_id']}/view?usp=sharing"
                st.markdown(f"**View/Download on Google Drive:** [Click Here]({drive_link})")
            else:
                st.warning("Report not available on Google Drive. It might have failed to upload or was generated before this feature was enabled.")
            
            # Option to re-generate and download if drive file is missing or for local access
            if st.button("Download as DOCX (Re-generate if needed)", key=f"download_report_{selected_report_id}"):
                with st.spinner("Re-generating report for download..."):
                    jd_text_for_regen = "" # You might need to fetch this from a stored JD or re-upload
                    cv_texts_for_regen = [] # Same here
                    
                    # For simplicity, we'll use a placeholder for actual content and rely on saved report_data
                    # In a real app, you'd retrieve JD/CV content from storage or re-process.
                    # For now, we'll just reconstruct the DOCX based on the saved report_data
                    report_buffer_regen = create_comparative_docx_report(
                        "Job Description text placeholder", # This text is not used in create_comparative_docx_report currently, only for function signature
                        ["CV text placeholder"], # This text is not used
                        selected_report['raw_data'], # Use the raw_data dictionary for docx generation
                        selected_report['raw_data'].get('candidate_evaluations', []),
                        selected_report['raw_data'].get('criteria_comparison_data', {}),
                        selected_report['raw_data'].get('general_and_shortlist_data', {})
                    )
                    st.download_button(
                        label="Download Re-generated Report",
                        data=report_buffer_regen.getvalue(),
                        file_name=f"{selected_report['generated_by_username']}_{selected_report['jd_filename'].replace('.pdf', '').replace('.docx', '')}_Analysis_{selected_report['timestamp'].replace(' ', '_').replace(':', '-')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        key=f"download_regen_button_{selected_report_id}"
                    )
                    st.success("Report re-generated and ready for download.")
            
            if st.session_state['is_admin']:
                if st.button("Delete Report (Admin Only)", key=f"delete_report_{selected_report_id}"):
                    # Added a confirmation step for deletion
                    if st.warning(f"Are you absolutely sure you want to delete {selected_report_id}? This action is irreversible."): # Changed message for clarity
                        if st.button("Confirm Deletion", key=f"confirm_delete_{selected_report_id}"): # Changed key for clarity
                            try:
                                # Delete from Firestore
                                db.collection('reports').document(selected_report_id).delete()
                                # Optionally: Delete from Google Drive too if drive_file_id exists
                                if selected_report['drive_file_id']:
                                    try:
                                        drive_service.files().delete(fileId=selected_report['drive_file_id']).execute()
                                        st.success("Report deleted from Google Drive.")
                                    except Exception as e:
                                        st.warning(f"Could not delete file from Google Drive: {e}. It might have been moved or already deleted.")
                                st.success("Report deleted successfully from database.")
                                st.rerun() # Refresh the page to show updated list
                            except Exception as e:
                                st.error(f"Error deleting report: {e}")

    except Exception as e:
        st.error(f"Error fetching reports: {e}")

def manage_users_page():
    if not st.session_state['is_admin']:
        st.error("Access Denied: You must be an admin to manage users.")
        return

    st.markdown("<h1 style='text-align: center; color: #4CAF50;'>SSO Consultants AI Recruitment Tool</h1>", unsafe_allow_html=True)
    st.subheader("Manage Users")

    # --- Invite New Member Section ---
    st.markdown("---")
    st.subheader("Invite New Member")
    with st.form("invite_new_member_form"):
        new_member_email = st.text_input("New Member's Email", key="invite_email")
        temp_password = st.text_input("Temporary Password", type="password", key="invite_temp_password")
        is_admin_invite = st.checkbox("Grant Admin Privileges", key="invite_is_admin_checkbox")
        
        confirm_admin_invite = True # Default to True if not admin invite
        if is_admin_invite:
            st.warning("Warning: Granting admin privileges provides full access to user management and report generation, including deleting reports and users.")
            confirm_admin_invite = st.checkbox("I understand and confirm to grant admin privileges to this user.", key="confirm_admin_privileges")
        
        submit_invite = st.form_submit_button("Invite Member")

        if submit_invite:
            if not new_member_email or not temp_password:
                st.error("Email and Temporary Password are required.")
            elif len(temp_password) < 6:
                st.error("Temporary Password must be at least 6 characters long.")
            elif is_admin_invite and not confirm_admin_invite:
                st.error("Please confirm granting admin privileges.")
            else:
                create_user(new_member_email, temp_password, is_admin_invite)
                st.rerun() # Refresh page to show updated user list and clear form

    st.markdown("---")
    st.write("Here you can view, activate/deactivate, and delete existing user accounts.")

    users_ref = db.collection('users')
    try:
        users_docs = users_ref.stream()
        users = []
        for doc in users_docs:
            user_data = doc.to_dict()
            try:
                firebase_user = auth.get_user(doc.id) # Fetch current status from Firebase Auth
                disabled_status = firebase_user.disabled
            except Exception:
                disabled_status = True # Assume disabled if user not found in Auth (e.g., deleted manually)

            users.append({
                'uid': doc.id,
                'email': user_data.get('email', 'N/A'),
                'username': user_data.get('username', 'Not Set'),
                'is_admin': user_data.get('is_admin', False),
                'disabled': disabled_status
            })
        
        # Filter out the current logged-in admin from the list of selectable users
        display_users = [u for u in users if u['uid'] != st.session_state['user_uid']]

        st.dataframe(
            display_users,
            column_order=["email", "username", "is_admin", "disabled"],
            hide_index=True,
            column_config={
                "uid": None, # Hide UID in table
                "email": "Email",
                "username": "Display Name",
                "is_admin": "Admin",
                "disabled": "Disabled"
            },
            use_container_width=True
        )

        st.markdown("---")
        st.subheader("Actions on Users")

        if display_users:
            selected_user_email = st.selectbox("Select a user to manage:", options=[u['email'] for u in display_users], key="user_selector")
            selected_user = next((u for u in display_users if u['email'] == selected_user_email), None)
            
            if selected_user:
                st.write(f"**Selected User:** {selected_user['username']} ({selected_user['email']})")
                st.write(f"Admin: {selected_user['is_admin']}, Disabled: {selected_user['disabled']}")

                if st.button(f"Toggle Admin Status ({'Revoke' if selected_user['is_admin'] else 'Grant'})", key=f"toggle_admin_{selected_user['uid']}"):
                    try:
                        new_admin_status = not selected_user['is_admin']
                        db.collection('users').document(selected_user['uid']).update({'is_admin': new_admin_status})
                        st.success(f"Admin status for {selected_user_email} changed to {new_admin_status}.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error toggling admin status: {e}")
                
                if st.button(f"Toggle Account Status ({'Enable' if selected_user['disabled'] else 'Disable'})", key=f"toggle_disabled_{selected_user['uid']}"):
                    try:
                        new_disabled_status = not selected_user['disabled']
                        auth.update_user(selected_user['uid'], disabled=new_disabled_status)
                        st.success(f"Account status for {selected_user_email} changed to {'disabled' if new_disabled_status else 'enabled'}.")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error toggling account status: {e}")

                if st.button("Delete User", key=f"delete_user_{selected_user['uid']}"):
                    # Added a confirmation step for deletion
                    if st.warning(f"Are you absolutely sure you want to delete {selected_user_email}? This action is irreversible."):
                        if st.button("Confirm Deletion", key=f"confirm_delete_{selected_user['uid']}"):
                            try:
                                # Delete from Firebase Auth
                                auth.delete_user(selected_user['uid'])
                                # Delete from Firestore
                                db.collection('users').document(selected_user['uid']).delete()
                                st.success(f"User {selected_user_email} deleted successfully.")
                                st.rerun()
                            except exceptions.FirebaseError as e:
                                st.error(f"Error deleting user: {e.code}")
                            except Exception as e:
                                st.error(f"An unexpected error occurred: {e}")
            else:
                st.info("Select a user from the dropdown to manage.")
        else:
            st.info("No other users to manage. You are the only admin.")

    except Exception as e:
        st.error(f"Error fetching users: {e}")

# --- Main Application Logic ---
# --- Fixed Position Logo ---
st.markdown(
    """
    <img src="data:image/png;base64,{}" class="fixed-top-left-logo">
    """.format(base64.b64encode(open("sso_logo.png", "rb").read()).decode()),
    unsafe_allow_html=True
)


if not st.session_state['logged_in']:
    if st.session_state['login_mode'] == 'choose_role':
        display_login_form()
    elif st.session_state['login_mode'] == 'login_form':
        show_login_and_create_account_forms()
elif st.session_state['needs_username_setup']:
    setup_username_and_password_page()
else:
    # Sidebar for navigation
    with st.sidebar:
        st.image("sso_logo.png", use_container_width=True)
        st.markdown(f"**Welcome, {st.session_state.get('username', st.session_state['user_email'])}!**")
        if st.session_state['is_admin']:
            st.markdown("**(Admin User)**")
        st.write("---")

        if st.session_state['is_admin']:
            if st.button("Generate Report", key="nav_generate_admin"):
                st.session_state['current_admin_page'] = 'generate'
                st.rerun()
            if st.button("View All Reports", key="nav_reports_admin"):
                st.session_state['current_admin_page'] = 'reports'
                st.rerun()
            if st.button("Manage Users", key="nav_manage_users_admin"):
                st.session_state['current_admin_page'] = 'manage_users'
                st.rerun()
        else: # Regular user view (MODIFIED)
            if st.button("Generate Report", key="nav_generate_user"):
                st.session_state['current_admin_page'] = 'generate'
                st.rerun()
            # Removed "View My Reports" button for regular users as per request
        
        st.write("---")
        if st.button("Logout", key="logout_button"):
            logout_user()

    # --- Main Content Area ---
    if st.session_state['is_admin']:
        if st.session_state['current_admin_page'] == 'generate':
            generate_comparative_report_page()
        elif st.session_state['current_admin_page'] == 'reports':
            show_all_reports_page()
        elif st.session_state['current_admin_page'] == 'manage_users':
            manage_users_page()
    else: # Regular user view (MODIFIED)
        # Regular users can only access the generate report page
        generate_comparative_report_page()


# --- Custom FOOTER (Always visible at the bottom of the page) ---
st.markdown(
    """
    <div style="
        position: fixed;
        bottom: 0;
        left: 0;
        width: 100%;
        text-align: center;
        color: #4CAF50 !important; /* Green text for footer */
        padding: 10px;
        background-color: #FFFFFF; /* Match page background */
        font-size: 0.8em;
        border-top: 1px solid #EEEEEE; /* Light grey border at the top of the footer */
        z-index: 999; /* Ensure footer is above most content but below fixed logo */
    ">
        SSO Consultants AI Recruitment Tool © 2025 | All Rights Reserved.
    </div>
    """,
    unsafe_allow_html=True
)
